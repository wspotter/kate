"""
Document processing service for Kate LLM Client.

This service handles extraction and processing of text from various document formats
for use in RAG and knowledge base functionality.
"""

import asyncio
import hashlib
import mimetypes
import re
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple, Union
from dataclasses import dataclass
from io import BytesIO

from loguru import logger

# Import document processing libraries
try:
    import PyPDF2
    import pdfplumber
    from pdf2image import convert_from_bytes
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    
try:
    import html2text
    HAS_HTML2TEXT = True
except ImportError:
    HAS_HTML2TEXT = False
    
try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

from ..core.events import EventBus
from ..database.manager import DatabaseManager
from ..database.models import Document, DocumentChunk


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    title: str
    content_type: str
    file_size: int
    word_count: int
    page_count: Optional[int] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    language: Optional[str] = None
    encoding: Optional[str] = None


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    content: str
    metadata: DocumentMetadata
    chunks: List[str]
    doc_hash: str
    processing_time_ms: int


@dataclass
class ChunkingConfig:
    """Configuration for text chunking."""
    chunk_size: int = 1000  # Characters per chunk
    chunk_overlap: int = 200  # Character overlap between chunks
    respect_sentence_boundaries: bool = True
    min_chunk_size: int = 100  # Minimum chunk size
    max_chunk_size: int = 2000  # Maximum chunk size


class DocumentProcessor:
    """
    Service for processing documents into text and chunks for RAG.
    
    Supports multiple document formats with intelligent text extraction
    and chunking strategies optimized for semantic search.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown', 
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel',
        '.csv': 'text/csv',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.rtf': 'application/rtf'
    }
    
    def __init__(
        self,
        database_manager: DatabaseManager,
        event_bus: EventBus,
        chunking_config: Optional[ChunkingConfig] = None
    ):
        self.database_manager = database_manager
        self.event_bus = event_bus
        self.chunking_config = chunking_config or ChunkingConfig()
        
        # Logger
        self.logger = logger.bind(component="DocumentProcessor")
        
        # Check available processors
        self._check_dependencies()
        
    def _check_dependencies(self) -> None:
        """Check which document processing libraries are available."""
        missing_deps = []
        
        if not HAS_PDF:
            missing_deps.append("PDF processing (PyPDF2, pdfplumber)")
        if not HAS_DOCX:
            missing_deps.append("DOCX processing (python-docx)")
        if not HAS_HTML2TEXT:
            missing_deps.append("HTML processing (html2text)")
        if not HAS_EXCEL:
            missing_deps.append("Excel processing (openpyxl)")
        if not HAS_OCR:
            missing_deps.append("OCR processing (pytesseract, PIL)")
            
        if missing_deps:
            self.logger.warning(f"Missing optional dependencies: {', '.join(missing_deps)}")
        else:
            self.logger.info("All document processing dependencies available")
            
    async def process_file(
        self, 
        file_path: Union[str, Path],
        title: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process a document file into text and chunks.
        
        Args:
            file_path: Path to the document file
            title: Optional custom title (defaults to filename)
            
        Returns:
            Processed document with content and metadata
        """
        import time
        start_time = time.time()
        
        try:
            file_path = Path(file_path)
            self.logger.info(f"Processing document: {file_path}")
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
                
            # Read file content
            with open(file_path, 'rb') as f:
                file_content = f.read()
                
            # Process the content
            result = await self.process_content(
                content=file_content,
                filename=file_path.name,
                title=title
            )
            
            processing_time = int((time.time() - start_time) * 1000)
            result.processing_time_ms = processing_time
            
            self.logger.info(f"Successfully processed {file_path} in {processing_time}ms")
            return result
            
        except Exception as e:
            self.logger.error(f"Failed to process file {file_path}: {e}")
            raise
            
    async def process_content(
        self,
        content: bytes,
        filename: str,
        title: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process document content from bytes.
        
        Args:
            content: Document content as bytes
            filename: Original filename for type detection
            title: Optional custom title
            
        Returns:
            Processed document with content and metadata
        """
        try:
            # Detect content type
            content_type = self._detect_content_type(content, filename)
            
            # Extract text based on content type
            extracted_text = await self._extract_text(content, content_type, filename)
            
            # Clean the text
            cleaned_text = self._clean_text(extracted_text)
            
            # Generate document hash
            doc_hash = self._generate_hash(content)
            
            # Create chunks
            chunks = self._create_chunks(cleaned_text)
            
            # Extract metadata
            metadata = DocumentMetadata(
                title=title or Path(filename).stem,
                content_type=content_type,
                file_size=len(content),
                word_count=len(cleaned_text.split()),
                encoding='utf-8'
            )
            
            return ProcessedDocument(
                content=cleaned_text,
                metadata=metadata,
                chunks=chunks,
                doc_hash=doc_hash,
                processing_time_ms=0  # Will be set by caller
            )
            
        except Exception as e:
            self.logger.error(f"Failed to process content for {filename}: {e}")
            raise
            
    def _detect_content_type(self, content: bytes, filename: str) -> str:
        """Detect the content type of a document."""
        # Try by file extension first
        file_ext = Path(filename).suffix.lower()
        if file_ext in self.SUPPORTED_EXTENSIONS:
            return self.SUPPORTED_EXTENSIONS[file_ext]
            
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            return mime_type
            
        # Try to detect by content
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK'):  # ZIP-based formats (DOCX, XLSX)
            if b'word/' in content:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif b'xl/' in content:
                return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif content.startswith(b'<html') or content.startswith(b'<!DOCTYPE'):
            return 'text/html'
            
        # Default to plain text
        return 'text/plain'
        
    async def _extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """Extract text from document content based on type."""
        try:
            if content_type == 'application/pdf':
                return await self._extract_pdf_text(content)
            elif 'wordprocessingml' in content_type or content_type == 'application/msword':
                return await self._extract_docx_text(content)
            elif content_type == 'text/html':
                return await self._extract_html_text(content)
            elif 'spreadsheetml' in content_type or 'ms-excel' in content_type:
                return await self._extract_excel_text(content)
            elif content_type == 'text/csv':
                return await self._extract_csv_text(content)
            elif content_type in ['text/plain', 'text/markdown']:
                return await self._extract_plain_text(content)
            elif content_type == 'application/json':
                return await self._extract_json_text(content)
            else:
                # Try as plain text
                return await self._extract_plain_text(content)
                
        except Exception as e:
            self.logger.warning(f"Failed to extract text from {content_type}: {e}")
            # Fallback to plain text
            return await self._extract_plain_text(content)
            
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content."""
        if not HAS_PDF:
            raise RuntimeError("PDF processing not available")
            
        text_parts = []
        
        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        
        except Exception as e:
            self.logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        
            except Exception as e2:
                self.logger.warning(f"PyPDF2 also failed: {e2}")
                
                # Last resort: OCR if available
                if HAS_OCR:
                    return await self._extract_pdf_ocr(content)
                    
        return '\n\n'.join(text_parts)
        
    async def _extract_pdf_ocr(self, content: bytes) -> str:
        """Extract text from PDF using OCR."""
        if not HAS_OCR:
            raise RuntimeError("OCR not available")
            
        try:
            # Convert PDF to images
            images = convert_from_bytes(content)
            
            text_parts = []
            for image in images:
                # OCR the image
                text = pytesseract.image_to_string(image)
                if text.strip():
                    text_parts.append(text)
                    
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"OCR extraction failed: {e}")
            return ""
            
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content."""
        if not HAS_DOCX:
            raise RuntimeError("DOCX processing not available")
            
        try:
            doc = DocxDocument(BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return ""
            
    async def _extract_html_text(self, content: bytes) -> str:
        """Extract text from HTML content."""
        if not HAS_HTML2TEXT:
            # Simple HTML stripping fallback
            import html
            text = content.decode('utf-8', errors='ignore')
            text = re.sub(r'<[^>]+>', '', text)
            return html.unescape(text)
            
        try:
            h = html2text.HTML2Text()
            h.ignore_links = True
            h.ignore_images = True
            text = h.handle(content.decode('utf-8', errors='ignore'))
            return text
            
        except Exception as e:
            self.logger.error(f"HTML extraction failed: {e}")
            return ""
            
    async def _extract_excel_text(self, content: bytes) -> str:
        """Extract text from Excel content."""
        if not HAS_EXCEL:
            raise RuntimeError("Excel processing not available")
            
        try:
            workbook = openpyxl.load_workbook(BytesIO(content), data_only=True)
            text_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        sheet_text.append('\t'.join(row_text))
                        
                if sheet_text:
                    text_parts.append(f"Sheet: {sheet_name}\n" + '\n'.join(sheet_text))
                    
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Excel extraction failed: {e}")
            return ""
            
    async def _extract_csv_text(self, content: bytes) -> str:
        """Extract text from CSV content."""
        try:
            import csv
            import io
            
            text = content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(io.StringIO(text))
            
            rows = []
            for row in csv_reader:
                if row:  # Skip empty rows
                    rows.append('\t'.join(row))
                    
            return '\n'.join(rows)
            
        except Exception as e:
            self.logger.error(f"CSV extraction failed: {e}")
            return ""
            
    async def _extract_plain_text(self, content: bytes) -> str:
        """Extract text from plain text content."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'ascii']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
                    
            # Fallback with error handling
            return content.decode('utf-8', errors='replace')
            
        except Exception as e:
            self.logger.error(f"Plain text extraction failed: {e}")
            return ""
            
    async def _extract_json_text(self, content: bytes) -> str:
        """Extract text from JSON content."""
        try:
            import json
            
            data = json.loads(content.decode('utf-8'))
            
            def extract_text_from_json(obj, path=""):
                """Recursively extract text values from JSON."""
                texts = []
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        new_path = f"{path}.{key}" if path else key
                        texts.extend(extract_text_from_json(value, new_path))
                elif isinstance(obj, list):
                    for i, value in enumerate(obj):
                        new_path = f"{path}[{i}]"
                        texts.extend(extract_text_from_json(value, new_path))
                elif isinstance(obj, str) and obj.strip():
                    texts.append(f"{path}: {obj}")
                elif obj is not None:
                    texts.append(f"{path}: {str(obj)}")
                    
                return texts
                
            text_parts = extract_text_from_json(data)
            return '\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"JSON extraction failed: {e}")
            return ""
            
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
            
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Trim whitespace
        text = text.strip()
        
        return text
        
    def _generate_hash(self, content: bytes) -> str:
        """Generate SHA-256 hash of document content."""
        return hashlib.sha256(content).hexdigest()
        
    def _create_chunks(self, text: str) -> List[str]:
        """Create text chunks for RAG processing."""
        if not text:
            return []
            
        chunks = []
        
        if self.chunking_config.respect_sentence_boundaries:
            chunks = self._chunk_by_sentences(text)
        else:
            chunks = self._chunk_by_size(text)
            
        # Filter chunks by size
        filtered_chunks = []
        for chunk in chunks:
            if len(chunk) >= self.chunking_config.min_chunk_size:
                if len(chunk) > self.chunking_config.max_chunk_size:
                    # Split large chunks further
                    sub_chunks = self._chunk_by_size(
                        chunk, 
                        self.chunking_config.max_chunk_size
                    )
                    filtered_chunks.extend(sub_chunks)
                else:
                    filtered_chunks.append(chunk)
                    
        return filtered_chunks
        
    def _chunk_by_sentences(self, text: str) -> List[str]:
        """Chunk text respecting sentence boundaries."""
        # Simple sentence splitting (can be improved with NLP libraries)
        sentences = re.split(r'[.!?]+\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            potential_chunk = f"{current_chunk} {sentence}".strip()
            
            if len(potential_chunk) <= self.chunking_config.chunk_size:
                current_chunk = potential_chunk
            else:
                # Start new chunk
                if current_chunk:
                    chunks.append(current_chunk)
                    
                # Handle overlap
                if self.chunking_config.chunk_overlap > 0 and chunks:
                    overlap_text = current_chunk[-self.chunking_config.chunk_overlap:]
                    current_chunk = f"{overlap_text} {sentence}".strip()
                else:
                    current_chunk = sentence
                    
        # Add final chunk
        if current_chunk:
            chunks.append(current_chunk)
            
        return chunks
        
    def _chunk_by_size(self, text: str, chunk_size: Optional[int] = None) -> List[str]:
        """Chunk text by character size."""
        if chunk_size is None:
            chunk_size = self.chunking_config.chunk_size
            
        chunks = []
        overlap = self.chunking_config.chunk_overlap
        
        for i in range(0, len(text), chunk_size - overlap):
            chunk = text[i:i + chunk_size]
            if chunk.strip():
                chunks.append(chunk.strip())
                
        return chunks
        
    async def store_document(self, processed_doc: ProcessedDocument, file_attachment_id: Optional[str] = None) -> Document:
        """
        Store a processed document in the database.
        
        Args:
            processed_doc: Processed document to store
            file_attachment_id: Optional file attachment ID
            
        Returns:
            Stored document record
        """
        try:
            # Create document record
            document = Document(
                title=processed_doc.metadata.title,
                content=processed_doc.content,
                content_type=processed_doc.metadata.content_type,
                file_attachment_id=file_attachment_id,
                doc_hash=processed_doc.doc_hash,
                word_count=processed_doc.metadata.word_count,
                extra_data={
                    "page_count": processed_doc.metadata.page_count,
                    "author": processed_doc.metadata.author,
                    "created_date": processed_doc.metadata.created_date,
                    "modified_date": processed_doc.metadata.modified_date,
                    "language": processed_doc.metadata.language,
                    "encoding": processed_doc.metadata.encoding,
                    "processing_time_ms": processed_doc.processing_time_ms
                }
            )
            
            # Store document
            async with self.database_manager.get_session() as session:
                session.add(document)
                await session.flush()  # Get the ID
                
                # Store chunks
                for i, chunk_text in enumerate(processed_doc.chunks):
                    chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk_text,
                        word_count=len(chunk_text.split())
                    )
                    session.add(chunk)
                    
                await session.commit()
                await session.refresh(document)
                
            self.logger.info(f"Stored document: {document.title} with {len(processed_doc.chunks)} chunks")
            return document
            
        except Exception as e:
            self.logger.error(f"Failed to store document: {e}")
            raise
            
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.SUPPORTED_EXTENSIONS.keys())
        
    def is_supported_format(self, filename: str) -> bool:
        """Check if a file format is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS